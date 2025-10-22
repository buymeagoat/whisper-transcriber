"""
Service layer for T034 Multi-Format Export System.
Provides comprehensive export functionality with format conversion and template processing.
"""

import json
import os
import zipfile
import tempfile
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
import logging

from sqlalchemy.orm import Session
from sqlalchemy import and_, or_, desc, asc, func

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from ..models.export_system import (
    ExportTemplate, ExportJob, BatchExport, ExportHistory, ExportFormatConfig,
    ExportFormat, ExportStatus, BatchExportStatus, TemplateType
)
from ..models import Job

logger = logging.getLogger(__name__)


class ExportFormatService:
    """Service for managing export format configurations and capabilities."""
    
    @staticmethod
    def get_available_formats(db: Session) -> List[Dict[str, Any]]:
        """Get all available export formats with their configurations."""
        configs = db.query(ExportFormatConfig).filter(ExportFormatConfig.is_enabled == True).all()
        
        formats = []
        for config in configs:
            formats.append({
                "format": config.format,
                "display_name": config.display_name,
                "file_extension": config.file_extension,
                "mime_type": config.mime_type,
                "supports_timestamps": config.supports_timestamps,
                "supports_styling": config.supports_styling,
                "supports_metadata": config.supports_metadata,
                "max_file_size_mb": config.max_file_size_mb,
                "default_config": config.default_config
            })
        
        return formats
    
    @staticmethod
    def get_format_config(db: Session, format_name: str) -> Optional[ExportFormatConfig]:
        """Get configuration for a specific export format."""
        return db.query(ExportFormatConfig).filter(
            and_(
                ExportFormatConfig.format == format_name,
                ExportFormatConfig.is_enabled == True
            )
        ).first()
    
    @staticmethod
    def validate_format_config(format_name: str, config: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """Validate export configuration against format requirements."""
        errors = []
        
        if format_name not in [fmt.value for fmt in ExportFormat]:
            errors.append(f"Unsupported export format: {format_name}")
            return False, errors
        
        # Format-specific validation
        if format_name == ExportFormat.SRT.value:
            if not isinstance(config.get("include_timestamps", True), bool):
                errors.append("SRT format requires boolean 'include_timestamps' setting")
        
        elif format_name == ExportFormat.PDF.value:
            if not PDF_AVAILABLE:
                errors.append("PDF export not available - reportlab library not installed")
        
        elif format_name == ExportFormat.DOCX.value:
            if not DOCX_AVAILABLE:
                errors.append("DOCX export not available - python-docx library not installed")
        
        return len(errors) == 0, errors


class ExportTemplateService:
    """Service for managing export templates and customization."""
    
    @staticmethod
    def create_template(
        db: Session,
        name: str,
        template_type: str,
        supported_formats: List[str],
        template_config: Dict[str, Any],
        styling_config: Optional[Dict[str, Any]] = None,
        layout_config: Optional[Dict[str, Any]] = None,
        created_by: Optional[str] = None,
        description: Optional[str] = None
    ) -> ExportTemplate:
        """Create a new export template."""
        
        template = ExportTemplate(
            name=name,
            description=description,
            template_type=template_type,
            supported_formats=supported_formats,
            template_config=template_config,
            styling_config=styling_config or {},
            layout_config=layout_config or {},
            created_by=created_by,
            is_system_template=False
        )
        
        db.add(template)
        db.commit()
        db.refresh(template)
        
        logger.info(f"Created export template: {name} (ID: {template.id})")
        return template
    
    @staticmethod
    def get_templates_for_format(db: Session, format_name: str) -> List[ExportTemplate]:
        """Get all available templates for a specific export format."""
        return db.query(ExportTemplate).filter(
            and_(
                ExportTemplate.is_active == True,
                ExportTemplate.supported_formats.contains([format_name])
            )
        ).order_by(ExportTemplate.is_system_template.desc(), ExportTemplate.name).all()
    
    @staticmethod
    def get_system_templates(db: Session) -> List[ExportTemplate]:
        """Get all system default templates."""
        return db.query(ExportTemplate).filter(
            and_(
                ExportTemplate.is_system_template == True,
                ExportTemplate.is_active == True
            )
        ).order_by(ExportTemplate.name).all()
    
    @staticmethod
    def create_default_templates(db: Session) -> None:
        """Create default system templates for all export formats."""
        
        # Default SRT template
        if not db.query(ExportTemplate).filter(
            and_(ExportTemplate.name == "Standard SRT", ExportTemplate.is_system_template == True)
        ).first():
            ExportTemplateService.create_system_template(
                db,
                name="Standard SRT",
                template_type=TemplateType.SUBTITLE.value,
                supported_formats=[ExportFormat.SRT.value],
                template_config={
                    "include_timestamps": True,
                    "timestamp_format": "HH:MM:SS,mmm",
                    "line_breaks": "\\n",
                    "max_line_length": 80,
                    "speaker_labels": False
                }
            )
        
        # Default VTT template
        if not db.query(ExportTemplate).filter(
            and_(ExportTemplate.name == "Standard WebVTT", ExportTemplate.is_system_template == True)
        ).first():
            ExportTemplateService.create_system_template(
                db,
                name="Standard WebVTT",
                template_type=TemplateType.SUBTITLE.value,
                supported_formats=[ExportFormat.VTT.value],
                template_config={
                    "include_timestamps": True,
                    "timestamp_format": "HH:MM:SS.mmm",
                    "line_breaks": "\\n",
                    "max_line_length": 80,
                    "speaker_labels": True,
                    "include_header": True,
                    "header_text": "WEBVTT"
                }
            )
        
        # Default DOCX template
        if not db.query(ExportTemplate).filter(
            and_(ExportTemplate.name == "Standard Document", ExportTemplate.is_system_template == True)
        ).first():
            ExportTemplateService.create_system_template(
                db,
                name="Standard Document",
                template_type=TemplateType.DOCUMENT.value,
                supported_formats=[ExportFormat.DOCX.value, ExportFormat.PDF.value],
                template_config={
                    "include_timestamps": True,
                    "include_metadata": True,
                    "paragraph_spacing": 1.15,
                    "show_speaker_labels": True,
                    "timestamp_in_margin": False
                },
                styling_config={
                    "font_family": "Calibri",
                    "font_size": 11,
                    "line_spacing": 1.15,
                    "margin_inches": 1.0
                }
            )
        
        # Default JSON template
        if not db.query(ExportTemplate).filter(
            and_(ExportTemplate.name == "Structured JSON", ExportTemplate.is_system_template == True)
        ).first():
            ExportTemplateService.create_system_template(
                db,
                name="Structured JSON",
                template_type=TemplateType.STRUCTURED.value,
                supported_formats=[ExportFormat.JSON.value],
                template_config={
                    "include_metadata": True,
                    "include_timestamps": True,
                    "segment_level": "sentence",
                    "include_confidence": True,
                    "include_speaker_info": True,
                    "pretty_print": True
                }
            )
    
    @staticmethod
    def create_system_template(db: Session, **kwargs) -> ExportTemplate:
        """Create a system template with is_system_template=True."""
        kwargs["is_system_template"] = True
        kwargs["created_by"] = "system"
        
        template = ExportTemplate(**kwargs)
        db.add(template)
        db.commit()
        db.refresh(template)
        return template


class SRTExportService:
    """Service for SRT subtitle format export."""
    
    @staticmethod
    def generate_srt(transcript_data: Dict[str, Any], config: Dict[str, Any]) -> str:
        """Generate SRT format subtitle file content."""
        
        segments = transcript_data.get("segments", [])
        if not segments:
            # Fallback to full text with estimated timestamps
            return SRTExportService._generate_srt_from_text(
                transcript_data.get("text", ""), config
            )
        
        srt_content = []
        
        for i, segment in enumerate(segments, 1):
            # Format timestamps
            start_time = SRTExportService._format_srt_timestamp(segment.get("start", 0))
            end_time = SRTExportService._format_srt_timestamp(segment.get("end", 0))
            
            # Format text
            text = segment.get("text", "").strip()
            if config.get("max_line_length", 80):
                text = SRTExportService._wrap_text(text, config["max_line_length"])
            
            # Add speaker labels if enabled
            if config.get("speaker_labels", False) and segment.get("speaker"):
                text = f"[{segment['speaker']}] {text}"
            
            # Add SRT entry
            srt_content.append(f"{i}")
            srt_content.append(f"{start_time} --> {end_time}")
            srt_content.append(text)
            srt_content.append("")  # Empty line between entries
        
        return "\n".join(srt_content)
    
    @staticmethod
    def _format_srt_timestamp(seconds: float) -> str:
        """Format seconds to SRT timestamp format (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"
    
    @staticmethod
    def _wrap_text(text: str, max_length: int) -> str:
        """Wrap text to maximum line length."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            word_length = len(word)
            if current_length + word_length + len(current_line) > max_length and current_line:
                lines.append(" ".join(current_line))
                current_line = [word]
                current_length = word_length
            else:
                current_line.append(word)
                current_length += word_length
        
        if current_line:
            lines.append(" ".join(current_line))
        
        return "\n".join(lines)
    
    @staticmethod
    def _generate_srt_from_text(text: str, config: Dict[str, Any]) -> str:
        """Generate SRT from plain text with estimated timestamps."""
        words_per_minute = config.get("words_per_minute", 150)
        max_chars_per_subtitle = config.get("max_line_length", 80) * 2
        
        sentences = text.split(". ")
        srt_content = []
        current_time = 0.0
        
        for i, sentence in enumerate(sentences, 1):
            if not sentence.strip():
                continue
                
            # Estimate duration based on character count
            char_count = len(sentence)
            duration = max(2.0, (char_count / max_chars_per_subtitle) * 4.0)  # Min 2 seconds
            
            start_time = SRTExportService._format_srt_timestamp(current_time)
            end_time = SRTExportService._format_srt_timestamp(current_time + duration)
            
            # Format text
            text_content = sentence.strip()
            if not text_content.endswith("."):
                text_content += "."
            
            if config.get("max_line_length", 80):
                text_content = SRTExportService._wrap_text(text_content, config["max_line_length"])
            
            srt_content.append(f"{i}")
            srt_content.append(f"{start_time} --> {end_time}")
            srt_content.append(text_content)
            srt_content.append("")
            
            current_time += duration + 0.5  # Small gap between subtitles
        
        return "\n".join(srt_content)


class VTTExportService:
    """Service for WebVTT subtitle format export."""
    
    @staticmethod
    def generate_vtt(transcript_data: Dict[str, Any], config: Dict[str, Any]) -> str:
        """Generate WebVTT format subtitle file content."""
        
        vtt_content = []
        
        # Add header
        if config.get("include_header", True):
            header = config.get("header_text", "WEBVTT")
            vtt_content.append(header)
            vtt_content.append("")
        
        segments = transcript_data.get("segments", [])
        if not segments:
            # Fallback to full text
            text_content = VTTExportService._generate_vtt_from_text(
                transcript_data.get("text", ""), config
            )
            vtt_content.append(text_content)
        else:
            for segment in segments:
                # Format timestamps
                start_time = VTTExportService._format_vtt_timestamp(segment.get("start", 0))
                end_time = VTTExportService._format_vtt_timestamp(segment.get("end", 0))
                
                # Format text
                text = segment.get("text", "").strip()
                
                # Add speaker labels if enabled
                if config.get("speaker_labels", True) and segment.get("speaker"):
                    text = f"<v {segment['speaker']}>{text}</v>"
                
                # Add line breaks for long text
                if config.get("max_line_length", 80):
                    text = VTTExportService._wrap_text_vtt(text, config["max_line_length"])
                
                vtt_content.append(f"{start_time} --> {end_time}")
                vtt_content.append(text)
                vtt_content.append("")
        
        return "\n".join(vtt_content)
    
    @staticmethod
    def _format_vtt_timestamp(seconds: float) -> str:
        """Format seconds to VTT timestamp format (HH:MM:SS.mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millisecs:03d}"
    
    @staticmethod
    def _wrap_text_vtt(text: str, max_length: int) -> str:
        """Wrap text for VTT with proper line breaks."""
        # Preserve speaker tags if present
        if text.startswith("<v ") and "</v>" in text:
            parts = text.split(">", 1)
            if len(parts) == 2:
                speaker_tag = parts[0] + ">"
                content = parts[1].replace("</v>", "")
                wrapped_content = SRTExportService._wrap_text(content, max_length)
                return speaker_tag + wrapped_content + "</v>"
        
        return SRTExportService._wrap_text(text, max_length)
    
    @staticmethod
    def _generate_vtt_from_text(text: str, config: Dict[str, Any]) -> str:
        """Generate VTT from plain text with estimated timestamps."""
        # Similar to SRT but with VTT formatting
        words_per_minute = config.get("words_per_minute", 150)
        max_chars_per_subtitle = config.get("max_line_length", 80) * 2
        
        sentences = text.split(". ")
        vtt_content = []
        current_time = 0.0
        
        for sentence in sentences:
            if not sentence.strip():
                continue
                
            char_count = len(sentence)
            duration = max(2.0, (char_count / max_chars_per_subtitle) * 4.0)
            
            start_time = VTTExportService._format_vtt_timestamp(current_time)
            end_time = VTTExportService._format_vtt_timestamp(current_time + duration)
            
            text_content = sentence.strip()
            if not text_content.endswith("."):
                text_content += "."
            
            if config.get("max_line_length", 80):
                text_content = VTTExportService._wrap_text_vtt(text_content, config["max_line_length"])
            
            vtt_content.append(f"{start_time} --> {end_time}")
            vtt_content.append(text_content)
            vtt_content.append("")
            
            current_time += duration + 0.5
        
        return "\n".join(vtt_content)


class DocumentExportService:
    """Service for document export (DOCX, PDF, TXT)."""
    
    @staticmethod
    def generate_docx(
        transcript_data: Dict[str, Any], 
        config: Dict[str, Any],
        styling_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate DOCX document from transcript data."""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available for DOCX export")
        
        doc = Document()
        
        # Apply styling
        style_config = styling_config or {}
        font_family = style_config.get("font_family", "Calibri")
        font_size = style_config.get("font_size", 11)
        
        # Document title and metadata
        if config.get("include_metadata", True):
            title = doc.add_heading("Transcript", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata table
            metadata = transcript_data.get("metadata", {})
            if metadata:
                doc.add_heading("Document Information", level=1)
                
                info_items = [
                    ("Original File", metadata.get("original_filename", "Unknown")),
                    ("Duration", f"{metadata.get('duration', 0):.1f} seconds"),
                    ("Model Used", metadata.get("model", "Unknown")),
                    ("Created", metadata.get("created_at", "Unknown")),
                ]
                
                for label, value in info_items:
                    p = doc.add_paragraph()
                    run1 = p.add_run(f"{label}: ")
                    run1.bold = True
                    run2 = p.add_run(str(value))
                    
                doc.add_page_break()
        
        # Add transcript content
        doc.add_heading("Transcript Content", level=1)
        
        segments = transcript_data.get("segments", [])
        if segments and config.get("include_timestamps", True):
            # Segmented transcript with timestamps
            for segment in segments:
                p = doc.add_paragraph()
                
                # Timestamp
                if not config.get("timestamp_in_margin", False):
                    time_str = DocumentExportService._format_document_timestamp(segment.get("start", 0))
                    run1 = p.add_run(f"[{time_str}] ")
                    run1.bold = True
                    run1.font.size = Pt(font_size - 1)
                
                # Speaker label
                if config.get("show_speaker_labels", True) and segment.get("speaker"):
                    run2 = p.add_run(f"{segment['speaker']}: ")
                    run2.bold = True
                
                # Text content
                text = segment.get("text", "").strip()
                run3 = p.add_run(text)
                
                # Apply font styling
                for run in p.runs:
                    run.font.name = font_family
                    if not run.font.size:
                        run.font.size = Pt(font_size)
        else:
            # Full text without segments
            text = transcript_data.get("text", "")
            paragraphs = text.split("\n")
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text.strip())
                    for run in p.runs:
                        run.font.name = font_family
                        run.font.size = Pt(font_size)
        
        # Save to bytes
        doc_bytes = DocumentExportService._docx_to_bytes(doc)
        return doc_bytes
    
    @staticmethod
    def generate_pdf(
        transcript_data: Dict[str, Any], 
        config: Dict[str, Any],
        styling_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """Generate PDF document from transcript data."""
        
        if not PDF_AVAILABLE:
            raise ImportError("reportlab library not available for PDF export")
        
        # Create temporary file for PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_path = temp_file.name
        
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                temp_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Build content
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph("Transcript", title_style))
            story.append(Spacer(1, 12))
            
            # Metadata
            if config.get("include_metadata", True):
                metadata = transcript_data.get("metadata", {})
                if metadata:
                    story.append(Paragraph("Document Information", styles['Heading2']))
                    
                    info_data = [
                        ["Original File", metadata.get("original_filename", "Unknown")],
                        ["Duration", f"{metadata.get('duration', 0):.1f} seconds"],
                        ["Model Used", metadata.get("model", "Unknown")],
                        ["Created", metadata.get("created_at", "Unknown")],
                    ]
                    
                    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
                    info_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                        ('BACKGROUND', (1, 0), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(info_table)
                    story.append(Spacer(1, 20))
            
            # Transcript content
            story.append(Paragraph("Transcript Content", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            segments = transcript_data.get("segments", [])
            if segments and config.get("include_timestamps", True):
                # Segmented transcript
                for segment in segments:
                    # Create paragraph text
                    text_parts = []
                    
                    # Timestamp
                    time_str = DocumentExportService._format_document_timestamp(segment.get("start", 0))
                    text_parts.append(f"<b>[{time_str}]</b> ")
                    
                    # Speaker
                    if config.get("show_speaker_labels", True) and segment.get("speaker"):
                        text_parts.append(f"<b>{segment['speaker']}:</b> ")
                    
                    # Text
                    text_parts.append(segment.get("text", "").strip())
                    
                    paragraph_text = "".join(text_parts)
                    story.append(Paragraph(paragraph_text, styles['Normal']))
                    story.append(Spacer(1, 6))
            else:
                # Full text
                text = transcript_data.get("text", "")
                paragraphs = text.split("\n")
                
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        story.append(Paragraph(paragraph_text.strip(), styles['Normal']))
                        story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            # Read PDF bytes
            with open(temp_path, "rb") as pdf_file:
                pdf_bytes = pdf_file.read()
            
            return pdf_bytes
        
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
    
    @staticmethod
    def generate_txt(transcript_data: Dict[str, Any], config: Dict[str, Any]) -> str:
        """Generate plain text format."""
        
        content_lines = []
        
        # Add metadata header if requested
        if config.get("include_metadata", True):
            metadata = transcript_data.get("metadata", {})
            if metadata:
                content_lines.append("TRANSCRIPT INFORMATION")
                content_lines.append("=" * 50)
                content_lines.append(f"Original File: {metadata.get('original_filename', 'Unknown')}")
                content_lines.append(f"Duration: {metadata.get('duration', 0):.1f} seconds")
                content_lines.append(f"Model Used: {metadata.get('model', 'Unknown')}")
                content_lines.append(f"Created: {metadata.get('created_at', 'Unknown')}")
                content_lines.append("")
                content_lines.append("TRANSCRIPT CONTENT")
                content_lines.append("=" * 50)
        
        # Add transcript content
        segments = transcript_data.get("segments", [])
        if segments and config.get("include_timestamps", True):
            # Segmented transcript with timestamps
            for segment in segments:
                line_parts = []
                
                # Timestamp
                time_str = DocumentExportService._format_document_timestamp(segment.get("start", 0))
                line_parts.append(f"[{time_str}]")
                
                # Speaker
                if config.get("show_speaker_labels", True) and segment.get("speaker"):
                    line_parts.append(f"{segment['speaker']}:")
                
                # Text
                line_parts.append(segment.get("text", "").strip())
                
                content_lines.append(" ".join(line_parts))
        else:
            # Full text without timestamps
            text = transcript_data.get("text", "")
            content_lines.append(text)
        
        return "\n".join(content_lines)
    
    @staticmethod
    def _format_document_timestamp(seconds: float) -> str:
        """Format timestamp for document display."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
    
    @staticmethod
    def _docx_to_bytes(doc) -> bytes:
        """Convert DOCX document to bytes."""
        with tempfile.NamedTemporaryFile() as temp_file:
            doc.save(temp_file.name)
            temp_file.seek(0)
            return temp_file.read()


class JSONExportService:
    """Service for JSON structured export."""
    
    @staticmethod
    def generate_json(transcript_data: Dict[str, Any], config: Dict[str, Any]) -> str:
        """Generate JSON format export."""
        
        # Build structured JSON
        export_data = {}
        
        # Metadata
        if config.get("include_metadata", True):
            export_data["metadata"] = transcript_data.get("metadata", {})
        
        # Transcript content
        segments = transcript_data.get("segments", [])
        if segments:
            export_data["segments"] = []
            
            for segment in segments:
                segment_data = {
                    "text": segment.get("text", "")
                }
                
                if config.get("include_timestamps", True):
                    segment_data["start"] = segment.get("start", 0)
                    segment_data["end"] = segment.get("end", 0)
                    segment_data["duration"] = segment.get("end", 0) - segment.get("start", 0)
                
                if config.get("include_confidence", True) and "confidence" in segment:
                    segment_data["confidence"] = segment["confidence"]
                
                if config.get("include_speaker_info", True) and "speaker" in segment:
                    segment_data["speaker"] = segment["speaker"]
                
                # Word-level data if available
                if config.get("segment_level") == "word" and "words" in segment:
                    segment_data["words"] = segment["words"]
                
                export_data["segments"].append(segment_data)
        
        # Full text
        export_data["full_text"] = transcript_data.get("text", "")
        
        # Processing information
        if config.get("include_metadata", True):
            export_data["processing_info"] = {
                "export_timestamp": datetime.utcnow().isoformat(),
                "export_format": "json",
                "segment_count": len(segments)
            }
        
        # Format JSON output
        if config.get("pretty_print", True):
            return json.dumps(export_data, indent=2, ensure_ascii=False)
        else:
            return json.dumps(export_data, ensure_ascii=False)


class ExportJobService:
    """Service for managing individual export jobs."""
    
    @staticmethod
    def create_export_job(
        db: Session,
        job_id: str,
        format: str,
        template_id: Optional[int] = None,
        custom_config: Optional[Dict[str, Any]] = None,
        created_by: Optional[str] = None,
        batch_export_id: Optional[int] = None
    ) -> ExportJob:
        """Create a new export job."""
        
        # Validate format
        if format not in [fmt.value for fmt in ExportFormat]:
            raise ValueError(f"Unsupported export format: {format}")
        
        # Set expiration (7 days from now)
        expires_at = datetime.utcnow() + timedelta(days=7)
        
        export_job = ExportJob(
            job_id=job_id,
            format=format,
            template_id=template_id,
            custom_config=custom_config or {},
            created_by=created_by,
            batch_export_id=batch_export_id,
            expires_at=expires_at
        )
        
        db.add(export_job)
        db.commit()
        db.refresh(export_job)
        
        logger.info(f"Created export job: {export_job.id} for job {job_id} in format {format}")
        return export_job
    
    @staticmethod
    def process_export_job(db: Session, export_job_id: int) -> bool:
        """Process an individual export job."""
        
        export_job = db.query(ExportJob).filter(ExportJob.id == export_job_id).first()
        if not export_job:
            logger.error(f"Export job {export_job_id} not found")
            return False
        
        try:
            # Update status to processing
            export_job.status = ExportStatus.PROCESSING.value
            export_job.processing_started_at = datetime.utcnow()
            export_job.progress_percentage = 0.0
            db.commit()
            
            # Get transcript data
            job = db.query(Job).filter(Job.id == export_job.job_id).first()
            if not job:
                raise ValueError(f"Source job {export_job.job_id} not found")
            
            # Prepare transcript data
            transcript_data = ExportJobService._prepare_transcript_data(job)
            export_job.progress_percentage = 25.0
            db.commit()
            
            # Get template configuration
            template_config, styling_config, layout_config = ExportJobService._get_template_config(
                db, export_job.template_id, export_job.custom_config
            )
            export_job.progress_percentage = 50.0
            db.commit()
            
            # Generate export content
            export_content = ExportJobService._generate_export_content(
                export_job.format, transcript_data, template_config, styling_config
            )
            export_job.progress_percentage = 75.0
            db.commit()
            
            # Save export file
            output_path, output_url = ExportJobService._save_export_file(
                export_job, export_content, job.original_filename or f"transcript_{job.id}"
            )
            
            # Update job completion
            export_job.status = ExportStatus.COMPLETED.value
            export_job.processing_completed_at = datetime.utcnow()
            export_job.progress_percentage = 100.0
            export_job.output_path = output_path
            export_job.output_url = output_url
            export_job.output_size_bytes = len(export_content) if isinstance(export_content, (str, bytes)) else None
            export_job.calculate_duration()
            
            db.commit()
            
            # Create history record
            ExportJobService._create_history_record(db, export_job, success=True)
            
            logger.info(f"Successfully processed export job {export_job_id}")
            return True
            
        except Exception as e:
            # Handle export failure
            export_job.status = ExportStatus.FAILED.value
            export_job.error_message = str(e)
            export_job.processing_completed_at = datetime.utcnow()
            export_job.calculate_duration()
            
            # Retry logic
            if export_job.retry_count < export_job.max_retries:
                export_job.retry_count += 1
                export_job.status = ExportStatus.PENDING.value
                export_job.error_message = None
                logger.warning(f"Export job {export_job_id} failed, retrying ({export_job.retry_count}/{export_job.max_retries})")
            else:
                logger.error(f"Export job {export_job_id} failed permanently: {str(e)}")
            
            db.commit()
            
            # Create failure history record
            ExportJobService._create_history_record(db, export_job, success=False, error=str(e))
            
            return False
    
    @staticmethod
    def _prepare_transcript_data(job: Job) -> Dict[str, Any]:
        """Prepare transcript data for export processing."""
        
        # Base transcript data structure
        transcript_data = {
            "text": job.result or "",
            "metadata": {
                "original_filename": job.original_filename,
                "duration": job.duration_seconds,
                "model": job.model,
                "created_at": job.created_at.isoformat() if job.created_at else None,
                "job_id": job.id,
                "language": getattr(job, "language", "auto"),
                "file_size": getattr(job, "file_size", None)
            }
        }
        
        # Add segments if available (from job.segments JSON field)
        segments_data = getattr(job, "segments", None)
        if segments_data:
            try:
                if isinstance(segments_data, str):
                    segments = json.loads(segments_data)
                else:
                    segments = segments_data
                
                if isinstance(segments, list) and segments:
                    transcript_data["segments"] = segments
            except (json.JSONDecodeError, TypeError):
                logger.warning(f"Could not parse segments data for job {job.id}")
        
        return transcript_data
    
    @staticmethod
    def _get_template_config(
        db: Session, 
        template_id: Optional[int], 
        custom_config: Optional[Dict[str, Any]]
    ) -> Tuple[Dict[str, Any], Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
        """Get template configuration for export processing."""
        
        template_config = {}
        styling_config = None
        layout_config = None
        
        if template_id:
            template = db.query(ExportTemplate).filter(ExportTemplate.id == template_id).first()
            if template:
                template_config = template.template_config or {}
                styling_config = template.styling_config
                layout_config = template.layout_config
        
        # Merge custom configuration
        if custom_config:
            template_config.update(custom_config)
        
        return template_config, styling_config, layout_config
    
    @staticmethod
    def _generate_export_content(
        format: str, 
        transcript_data: Dict[str, Any], 
        template_config: Dict[str, Any],
        styling_config: Optional[Dict[str, Any]] = None
    ) -> Union[str, bytes]:
        """Generate export content based on format."""
        
        if format == ExportFormat.SRT.value:
            return SRTExportService.generate_srt(transcript_data, template_config)
        
        elif format == ExportFormat.VTT.value:
            return VTTExportService.generate_vtt(transcript_data, template_config)
        
        elif format == ExportFormat.DOCX.value:
            return DocumentExportService.generate_docx(transcript_data, template_config, styling_config)
        
        elif format == ExportFormat.PDF.value:
            return DocumentExportService.generate_pdf(transcript_data, template_config, styling_config)
        
        elif format == ExportFormat.TXT.value:
            return DocumentExportService.generate_txt(transcript_data, template_config)
        
        elif format == ExportFormat.JSON.value:
            return JSONExportService.generate_json(transcript_data, template_config)
        
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    @staticmethod
    def _save_export_file(
        export_job: ExportJob, 
        content: Union[str, bytes], 
        base_filename: str
    ) -> Tuple[str, str]:
        """Save export content to file and return paths."""
        
        # Create export directory
        export_dir = Path("exports") / str(export_job.id)
        export_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename
        format_config = {
            ExportFormat.SRT.value: ".srt",
            ExportFormat.VTT.value: ".vtt",
            ExportFormat.DOCX.value: ".docx",
            ExportFormat.PDF.value: ".pdf",
            ExportFormat.TXT.value: ".txt",
            ExportFormat.JSON.value: ".json"
        }
        
        extension = format_config.get(export_job.format, ".txt")
        filename = f"{Path(base_filename).stem}_{export_job.format}{extension}"
        filepath = export_dir / filename
        
        # Save content
        if isinstance(content, str):
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
        else:
            with open(filepath, 'wb') as f:
                f.write(content)
        
        # Generate download URL (this would be handled by your file serving endpoint)
        download_url = f"/api/exports/{export_job.id}/download"
        
        export_job.output_filename = filename
        
        return str(filepath), download_url
    
    @staticmethod
    def _create_history_record(
        db: Session, 
        export_job: ExportJob, 
        success: bool, 
        error: Optional[str] = None
    ) -> None:
        """Create export history record."""
        
        history = ExportHistory(
            export_job_id=export_job.id,
            batch_export_id=export_job.batch_export_id,
            export_type="single",
            format=export_job.format,
            template_name=export_job.template.name if export_job.template else None,
            user_id=export_job.created_by,
            processing_time_seconds=export_job.processing_duration_seconds,
            output_size_bytes=export_job.output_size_bytes,
            success=success,
            error_type=type(Exception).__name__ if error else None,
            error_details=error
        )
        
        db.add(history)
        db.commit()


class BatchExportService:
    """Service for managing batch export operations."""
    
    @staticmethod
    def create_batch_export(
        db: Session,
        name: str,
        job_ids: List[str],
        export_format: str,
        template_id: Optional[int] = None,
        batch_config: Optional[Dict[str, Any]] = None,
        created_by: Optional[str] = None,
        description: Optional[str] = None
    ) -> BatchExport:
        """Create a new batch export operation."""
        
        # Validate format
        if export_format not in [fmt.value for fmt in ExportFormat]:
            raise ValueError(f"Unsupported export format: {export_format}")
        
        # Set expiration (14 days for batch exports)
        expires_at = datetime.utcnow() + timedelta(days=14)
        
        batch_export = BatchExport(
            name=name,
            description=description,
            export_format=export_format,
            template_id=template_id,
            batch_config=batch_config or {},
            job_ids=job_ids,
            total_jobs=len(job_ids),
            created_by=created_by,
            expires_at=expires_at
        )
        
        db.add(batch_export)
        db.commit()
        db.refresh(batch_export)
        
        # Create individual export jobs
        for job_id in job_ids:
            ExportJobService.create_export_job(
                db=db,
                job_id=job_id,
                format=export_format,
                template_id=template_id,
                custom_config=batch_config,
                created_by=created_by,
                batch_export_id=batch_export.id
            )
        
        logger.info(f"Created batch export: {batch_export.id} with {len(job_ids)} jobs")
        return batch_export
    
    @staticmethod
    def process_batch_export(db: Session, batch_export_id: int) -> bool:
        """Process a batch export by processing all individual export jobs."""
        
        batch_export = db.query(BatchExport).filter(BatchExport.id == batch_export_id).first()
        if not batch_export:
            logger.error(f"Batch export {batch_export_id} not found")
            return False
        
        try:
            # Update status to processing
            batch_export.status = BatchExportStatus.PROCESSING.value
            batch_export.started_at = datetime.utcnow()
            db.commit()
            
            # Process all export jobs
            export_jobs = db.query(ExportJob).filter(
                ExportJob.batch_export_id == batch_export_id
            ).all()
            
            successful_jobs = []
            failed_jobs = []
            
            for export_job in export_jobs:
                success = ExportJobService.process_export_job(db, export_job.id)
                if success:
                    successful_jobs.append(export_job)
                    batch_export.completed_jobs += 1
                else:
                    failed_jobs.append(export_job)
                    batch_export.failed_jobs += 1
                
                batch_export.update_progress()
                db.commit()
            
            # Create archive if any jobs succeeded
            if successful_jobs:
                archive_path = BatchExportService._create_batch_archive(batch_export, successful_jobs)
                batch_export.archive_path = archive_path
                batch_export.download_url = f"/api/batch-exports/{batch_export.id}/download"
                
                # Calculate archive size
                if os.path.exists(archive_path):
                    batch_export.archive_size_bytes = os.path.getsize(archive_path)
            
            # Determine final status
            if batch_export.failed_jobs == 0:
                batch_export.status = BatchExportStatus.COMPLETED.value
            elif batch_export.completed_jobs > 0:
                batch_export.status = BatchExportStatus.PARTIAL.value
                batch_export.partial_success = True
            else:
                batch_export.status = BatchExportStatus.FAILED.value
            
            batch_export.completed_at = datetime.utcnow()
            batch_export.calculate_duration()
            db.commit()
            
            # Create history record
            BatchExportService._create_batch_history_record(db, batch_export, len(successful_jobs) > 0)
            
            logger.info(f"Completed batch export {batch_export_id}: {len(successful_jobs)} succeeded, {len(failed_jobs)} failed")
            return len(successful_jobs) > 0
            
        except Exception as e:
            batch_export.status = BatchExportStatus.FAILED.value
            batch_export.error_message = str(e)
            batch_export.completed_at = datetime.utcnow()
            batch_export.calculate_duration()
            db.commit()
            
            logger.error(f"Batch export {batch_export_id} failed: {str(e)}")
            return False
    
    @staticmethod
    def _create_batch_archive(batch_export: BatchExport, successful_jobs: List[ExportJob]) -> str:
        """Create ZIP archive containing all successful exports."""
        
        # Create archive directory
        archive_dir = Path("exports") / "batches" / str(batch_export.id)
        archive_dir.mkdir(parents=True, exist_ok=True)
        
        # Archive filename
        archive_filename = f"{batch_export.name}_{batch_export.export_format}_batch.zip"
        archive_path = archive_dir / archive_filename
        
        # Create ZIP archive
        with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for export_job in successful_jobs:
                if export_job.output_path and os.path.exists(export_job.output_path):
                    # Add file to archive with appropriate name
                    arcname = export_job.output_filename or f"export_{export_job.id}"
                    zipf.write(export_job.output_path, arcname)
        
        batch_export.archive_filename = archive_filename
        return str(archive_path)
    
    @staticmethod
    def _create_batch_history_record(db: Session, batch_export: BatchExport, success: bool) -> None:
        """Create batch export history record."""
        
        history = ExportHistory(
            batch_export_id=batch_export.id,
            export_type="batch",
            format=batch_export.export_format,
            template_name=batch_export.template.name if batch_export.template else None,
            user_id=batch_export.created_by,
            processing_time_seconds=batch_export.processing_duration_seconds,
            output_size_bytes=batch_export.archive_size_bytes,
            job_count=batch_export.total_jobs,
            success=success,
            error_type="BatchProcessingError" if not success else None,
            error_details=batch_export.error_message if not success else None
        )
        
        db.add(history)
        db.commit()