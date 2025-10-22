"""
T022: Multi-Format Export System - Export Service
Comprehensive service for exporting transcripts to multiple formats
Supports SRT, VTT, DOCX, PDF, JSON with customizable templates
"""

import json
import io
import re
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union, BinaryIO
from enum import Enum
from dataclasses import dataclass, asdict
from pathlib import Path
import logging

# Third-party imports for document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Import models directly from models.py to avoid circular imports
import sys
if 'api.models' not in sys.modules:
    from api import models as api_models
    Job = api_models.Job
else:
    from api.models import Job

# TranscriptMetadata may not exist, handle gracefully
try:
    from api.models import TranscriptMetadata
except (ImportError, AttributeError):
    # Define a minimal placeholder for TranscriptMetadata
    class TranscriptMetadata:
        def __init__(self, **kwargs):
            for k, v in kwargs.items():
                setattr(self, k, v)

logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """Supported export formats"""
    SRT = "srt"
    VTT = "vtt"
    DOCX = "docx"
    PDF = "pdf"
    JSON = "json"
    TXT = "txt"


@dataclass
class ExportTemplate:
    """Template configuration for export formats"""
    name: str
    format: ExportFormat
    description: str
    include_timestamps: bool = True
    include_metadata: bool = True
    include_summary: bool = False
    include_keywords: bool = False
    custom_styling: Optional[Dict[str, Any]] = None
    timestamp_format: str = "HH:MM:SS"


@dataclass
class ExportOptions:
    """Options for export customization"""
    template: Optional[ExportTemplate] = None
    custom_filename: Optional[str] = None
    include_header: bool = True
    include_footer: bool = True
    word_wrap: int = 80
    font_size: int = 12
    font_family: str = "Arial"
    line_spacing: float = 1.5
    page_margins: Dict[str, float] = None
    
    def __post_init__(self):
        if self.page_margins is None:
            self.page_margins = {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}


@dataclass
class ExportResult:
    """Result of export operation"""
    success: bool
    format: ExportFormat
    filename: str
    file_size: int
    content: Optional[Union[str, bytes]] = None
    error: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class TranscriptExportService:
    """Service for exporting transcripts to multiple formats"""
    
    def __init__(self):
        self.default_templates = self._create_default_templates()
        self.export_handlers = {
            ExportFormat.SRT: self._export_srt,
            ExportFormat.VTT: self._export_vtt,
            ExportFormat.JSON: self._export_json,
            ExportFormat.TXT: self._export_txt,
        }
        
        # Add optional handlers if libraries are available
        if DOCX_AVAILABLE:
            self.export_handlers[ExportFormat.DOCX] = self._export_docx
        if PDF_AVAILABLE:
            self.export_handlers[ExportFormat.PDF] = self._export_pdf
    
    def _create_default_templates(self) -> Dict[str, ExportTemplate]:
        """Create default export templates"""
        return {
            "srt_standard": ExportTemplate(
                name="Standard SRT",
                format=ExportFormat.SRT,
                description="Standard SubRip format with timestamps",
                include_timestamps=True,
                include_metadata=False
            ),
            "vtt_web": ExportTemplate(
                name="Web VTT",
                format=ExportFormat.VTT,
                description="WebVTT format for web video players",
                include_timestamps=True,
                include_metadata=True
            ),
            "docx_report": ExportTemplate(
                name="Document Report",
                format=ExportFormat.DOCX,
                description="Professional document with metadata",
                include_timestamps=True,
                include_metadata=True,
                include_summary=True,
                include_keywords=True
            ),
            "pdf_transcript": ExportTemplate(
                name="PDF Transcript",
                format=ExportFormat.PDF,
                description="Formatted PDF document",
                include_timestamps=True,
                include_metadata=True,
                include_summary=True
            ),
            "json_full": ExportTemplate(
                name="Complete JSON",
                format=ExportFormat.JSON,
                description="Full transcript data in JSON format",
                include_timestamps=True,
                include_metadata=True,
                include_summary=True,
                include_keywords=True
            ),
            "txt_clean": ExportTemplate(
                name="Plain Text",
                format=ExportFormat.TXT,
                description="Clean text without timestamps",
                include_timestamps=False,
                include_metadata=False
            )
        }
    
    def get_available_formats(self) -> List[Dict[str, Any]]:
        """Get list of available export formats"""
        formats = []
        for fmt in ExportFormat:
            available = fmt in self.export_handlers
            formats.append({
                "format": fmt.value,
                "name": fmt.value.upper(),
                "available": available,
                "description": self._get_format_description(fmt),
                "requires": self._get_format_requirements(fmt)
            })
        return formats
    
    def _get_format_description(self, fmt: ExportFormat) -> str:
        """Get description for export format"""
        descriptions = {
            ExportFormat.SRT: "SubRip subtitle format (.srt)",
            ExportFormat.VTT: "WebVTT subtitle format (.vtt)",
            ExportFormat.DOCX: "Microsoft Word document (.docx)",
            ExportFormat.PDF: "Portable Document Format (.pdf)",
            ExportFormat.JSON: "JavaScript Object Notation (.json)",
            ExportFormat.TXT: "Plain text file (.txt)"
        }
        return descriptions.get(fmt, "Unknown format")
    
    def _get_format_requirements(self, fmt: ExportFormat) -> List[str]:
        """Get requirements for export format"""
        requirements = {
            ExportFormat.DOCX: ["python-docx"] if not DOCX_AVAILABLE else [],
            ExportFormat.PDF: ["reportlab"] if not PDF_AVAILABLE else [],
        }
        return requirements.get(fmt, [])
    
    def get_templates(self, format_filter: Optional[ExportFormat] = None) -> List[Dict[str, Any]]:
        """Get available export templates"""
        templates = []
        for template in self.default_templates.values():
            if format_filter is None or template.format == format_filter:
                template_dict = asdict(template)
                template_dict["format"] = template.format.value
                templates.append(template_dict)
        return templates
    
    def export_transcript(
        self,
        job: Job,
        format: ExportFormat,
        options: Optional[ExportOptions] = None
    ) -> ExportResult:
        """Export transcript to specified format"""
        try:
            # Validate format availability
            if format not in self.export_handlers:
                return ExportResult(
                    success=False,
                    format=format,
                    filename="",
                    file_size=0,
                    error=f"Export format {format.value} not available. Missing dependencies: {self._get_format_requirements(format)}"
                )
            
            # Use default options if none provided
            if options is None:
                options = ExportOptions()
            
            # Set template if not provided
            if options.template is None:
                options.template = self._get_default_template(format)
            
            # Generate filename
            filename = self._generate_filename(job, format, options.custom_filename)
            
            # Call appropriate export handler
            handler = self.export_handlers[format]
            content = handler(job, options)
            
            # Calculate file size
            if isinstance(content, str):
                file_size = len(content.encode('utf-8'))
            else:
                file_size = len(content)
            
            # Create metadata
            metadata = self._create_export_metadata(job, format, options)
            
            logger.info(f"Successfully exported transcript {job.id} to {format.value} format ({file_size} bytes)")
            
            return ExportResult(
                success=True,
                format=format,
                filename=filename,
                file_size=file_size,
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Failed to export transcript {job.id} to {format.value}: {str(e)}")
            return ExportResult(
                success=False,
                format=format,
                filename="",
                file_size=0,
                error=str(e)
            )
    
    def _get_default_template(self, format: ExportFormat) -> ExportTemplate:
        """Get default template for format"""
        format_templates = {
            ExportFormat.SRT: "srt_standard",
            ExportFormat.VTT: "vtt_web",
            ExportFormat.DOCX: "docx_report",
            ExportFormat.PDF: "pdf_transcript",
            ExportFormat.JSON: "json_full",
            ExportFormat.TXT: "txt_clean"
        }
        template_key = format_templates.get(format, "srt_standard")
        return self.default_templates[template_key]
    
    def _generate_filename(
        self,
        job: Job,
        format: ExportFormat,
        custom_filename: Optional[str] = None
    ) -> str:
        """Generate export filename"""
        if custom_filename:
            # Ensure proper extension
            if not custom_filename.endswith(f".{format.value}"):
                return f"{custom_filename}.{format.value}"
            return custom_filename
        
        # Use original filename as base
        base_name = Path(job.original_filename).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base_name}_transcript_{timestamp}.{format.value}"
    
    def _create_export_metadata(
        self,
        job: Job,
        format: ExportFormat,
        options: ExportOptions
    ) -> Dict[str, Any]:
        """Create export metadata"""
        return {
            "job_id": job.id,
            "original_filename": job.original_filename,
            "export_format": format.value,
            "export_timestamp": datetime.utcnow().isoformat(),
            "template_name": options.template.name if options.template else None,
            "options": {
                "include_timestamps": options.template.include_timestamps if options.template else True,
                "include_metadata": options.template.include_metadata if options.template else True,
                "font_size": options.font_size,
                "font_family": options.font_family
            },
            "transcript_metadata": {
                "duration": job.transcript_metadata.duration if job.transcript_metadata else None,
                "language": job.transcript_metadata.language if job.transcript_metadata else None,
                "model": job.transcript_metadata.model if job.transcript_metadata else None,
                "confidence_score": job.transcript_metadata.confidence_score if job.transcript_metadata else None,
                "word_count": job.transcript_metadata.word_count if job.transcript_metadata else None
            }
        }
    
    def _export_srt(self, job: Job, options: ExportOptions) -> str:
        """Export to SRT format"""
        if not job.transcript_content:
            raise ValueError("No transcript content available")
        
        # Parse transcript content to extract timing information
        # For now, create basic SRT with estimated timings
        lines = job.transcript_content.strip().split('\n')
        srt_content = []
        
        # Estimate timing based on words per minute (150 WPM average)
        words_per_minute = 150
        total_words = sum(len(line.split()) for line in lines if line.strip())
        
        if job.transcript_metadata and job.transcript_metadata.duration:
            duration_seconds = job.transcript_metadata.duration
        else:
            # Estimate duration based on word count
            duration_seconds = (total_words / words_per_minute) * 60
        
        current_time = 0.0
        for i, line in enumerate(lines, 1):
            if not line.strip():
                continue
                
            words_in_line = len(line.split())
            line_duration = (words_in_line / words_per_minute) * 60
            
            start_time = self._seconds_to_srt_time(current_time)
            end_time = self._seconds_to_srt_time(current_time + line_duration)
            
            srt_content.append(f"{i}")
            srt_content.append(f"{start_time} --> {end_time}")
            srt_content.append(line.strip())
            srt_content.append("")  # Empty line between subtitles
            
            current_time += line_duration
        
        return '\n'.join(srt_content)
    
    def _seconds_to_srt_time(self, seconds: float) -> str:
        """Convert seconds to SRT time format (HH:MM:SS,mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"
    
    def _export_vtt(self, job: Job, options: ExportOptions) -> str:
        """Export to VTT format"""
        if not job.transcript_content:
            raise ValueError("No transcript content available")
        
        vtt_content = ["WEBVTT"]
        
        # Add metadata if requested
        if options.template.include_metadata and job.transcript_metadata:
            vtt_content.append("")
            vtt_content.append("NOTE")
            vtt_content.append(f"Original file: {job.original_filename}")
            if job.transcript_metadata.language:
                vtt_content.append(f"Language: {job.transcript_metadata.language}")
            if job.transcript_metadata.model:
                vtt_content.append(f"Model: {job.transcript_metadata.model}")
        
        vtt_content.append("")
        
        # Parse transcript similar to SRT but with VTT format
        lines = job.transcript_content.strip().split('\n')
        words_per_minute = 150
        
        if job.transcript_metadata and job.transcript_metadata.duration:
            duration_seconds = job.transcript_metadata.duration
        else:
            total_words = sum(len(line.split()) for line in lines if line.strip())
            duration_seconds = (total_words / words_per_minute) * 60
        
        current_time = 0.0
        for line in lines:
            if not line.strip():
                continue
                
            words_in_line = len(line.split())
            line_duration = (words_in_line / words_per_minute) * 60
            
            start_time = self._seconds_to_vtt_time(current_time)
            end_time = self._seconds_to_vtt_time(current_time + line_duration)
            
            vtt_content.append(f"{start_time} --> {end_time}")
            vtt_content.append(line.strip())
            vtt_content.append("")
            
            current_time += line_duration
        
        return '\n'.join(vtt_content)
    
    def _seconds_to_vtt_time(self, seconds: float) -> str:
        """Convert seconds to VTT time format (HH:MM:SS.mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millisecs:03d}"
    
    def _export_json(self, job: Job, options: ExportOptions) -> str:
        """Export to JSON format"""
        export_data = {
            "job_id": job.id,
            "original_filename": job.original_filename,
            "created_at": job.created_at.isoformat() if job.created_at else None,
            "status": job.status,
            "transcript_content": job.transcript_content
        }
        
        # Add metadata if available and requested
        if options.template.include_metadata and job.transcript_metadata:
            metadata = job.transcript_metadata
            export_data["metadata"] = {
                "duration": metadata.duration,
                "language": metadata.language,
                "model": metadata.model,
                "confidence_score": metadata.confidence_score,
                "word_count": metadata.word_count,
                "file_size": metadata.file_size,
                "audio_format": metadata.audio_format,
                "sample_rate": metadata.sample_rate,
                "channels": metadata.channels
            }
            
            # Add optional fields
            if options.template.include_summary and metadata.summary:
                export_data["metadata"]["summary"] = metadata.summary
            
            if options.template.include_keywords and metadata.keywords:
                export_data["metadata"]["keywords"] = metadata.keywords
        
        # Add export metadata
        export_data["export_info"] = {
            "format": "json",
            "template": options.template.name,
            "exported_at": datetime.utcnow().isoformat(),
            "options": asdict(options.template)
        }
        
        return json.dumps(export_data, indent=2, ensure_ascii=False)
    
    def _export_txt(self, job: Job, options: ExportOptions) -> str:
        """Export to plain text format"""
        if not job.transcript_content:
            raise ValueError("No transcript content available")
        
        content = []
        
        # Add header if requested
        if options.include_header:
            content.append(f"Transcript: {job.original_filename}")
            if job.transcript_metadata and options.template.include_metadata:
                if job.transcript_metadata.language:
                    content.append(f"Language: {job.transcript_metadata.language}")
                if job.transcript_metadata.duration:
                    duration_str = self._format_duration(job.transcript_metadata.duration)
                    content.append(f"Duration: {duration_str}")
            content.append("-" * 50)
            content.append("")
        
        # Add transcript content
        lines = job.transcript_content.strip().split('\n')
        if options.word_wrap > 0:
            # Apply word wrapping
            wrapped_content = []
            for line in lines:
                if len(line) <= options.word_wrap:
                    wrapped_content.append(line)
                else:
                    # Simple word wrapping
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + " " + word) <= options.word_wrap:
                            current_line = current_line + " " + word if current_line else word
                        else:
                            wrapped_content.append(current_line)
                            current_line = word
                    if current_line:
                        wrapped_content.append(current_line)
            content.extend(wrapped_content)
        else:
            content.extend(lines)
        
        # Add footer if requested
        if options.include_footer:
            content.append("")
            content.append("-" * 50)
            content.append(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        return '\n'.join(content)
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration in human-readable format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
    
    # Optional document format exports (require additional libraries)
    
    def _export_docx(self, job: Job, options: ExportOptions) -> bytes:
        """Export to DOCX format (requires python-docx)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available")
        
        if not job.transcript_content:
            raise ValueError("No transcript content available")
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Transcript: {job.original_filename}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        if options.template.include_metadata and job.transcript_metadata:
            doc.add_heading("Metadata", level=1)
            metadata_table = doc.add_table(rows=0, cols=2)
            metadata_table.style = 'Light Shading Accent 1'
            
            metadata_items = [
                ("Original File", job.original_filename),
                ("Language", job.transcript_metadata.language or "Unknown"),
                ("Model", job.transcript_metadata.model or "Unknown"),
                ("Duration", self._format_duration(job.transcript_metadata.duration) if job.transcript_metadata.duration else "Unknown"),
                ("Confidence", f"{job.transcript_metadata.confidence_score:.1%}" if job.transcript_metadata.confidence_score else "Unknown"),
                ("Word Count", str(job.transcript_metadata.word_count) if job.transcript_metadata.word_count else "Unknown")
            ]
            
            for key, value in metadata_items:
                row = metadata_table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(value)
        
        # Add summary if available
        if (options.template.include_summary and 
            job.transcript_metadata and 
            job.transcript_metadata.summary):
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(job.transcript_metadata.summary)
        
        # Add keywords if available
        if (options.template.include_keywords and 
            job.transcript_metadata and 
            job.transcript_metadata.keywords):
            doc.add_heading("Keywords", level=1)
            keywords_text = ", ".join(job.transcript_metadata.keywords)
            doc.add_paragraph(keywords_text)
        
        # Add transcript content
        doc.add_heading("Transcript", level=1)
        
        # Split content into paragraphs
        lines = job.transcript_content.strip().split('\n')
        for line in lines:
            if line.strip():
                para = doc.add_paragraph(line.strip())
                # Apply formatting
                run = para.runs[0]
                run.font.name = options.font_family
                run.font.size = Pt(options.font_size)
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
    
    def _export_pdf(self, job: Job, options: ExportOptions) -> bytes:
        """Export to PDF format (requires reportlab)"""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab library not available")
        
        if not job.transcript_content:
            raise ValueError("No transcript content available")
        
        # Create PDF document
        pdf_io = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_io,
            pagesize=letter,
            rightMargin=options.page_margins["right"] * inch,
            leftMargin=options.page_margins["left"] * inch,
            topMargin=options.page_margins["top"] * inch,
            bottomMargin=options.page_margins["bottom"] * inch
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=options.font_size,
            fontName=options.font_family,
            leading=options.font_size * options.line_spacing
        )
        
        # Build document content
        story = []
        
        # Title
        story.append(Paragraph(f"Transcript: {job.original_filename}", title_style))
        story.append(Spacer(1, 12))
        
        # Metadata table
        if options.template.include_metadata and job.transcript_metadata:
            metadata_data = [
                ['Property', 'Value'],
                ['Original File', job.original_filename],
                ['Language', job.transcript_metadata.language or 'Unknown'],
                ['Model', job.transcript_metadata.model or 'Unknown']
            ]
            
            if job.transcript_metadata.duration:
                metadata_data.append(['Duration', self._format_duration(job.transcript_metadata.duration)])
            if job.transcript_metadata.confidence_score:
                metadata_data.append(['Confidence', f"{job.transcript_metadata.confidence_score:.1%}"])
            if job.transcript_metadata.word_count:
                metadata_data.append(['Word Count', str(job.transcript_metadata.word_count)])
            
            metadata_table = Table(metadata_data)
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(metadata_table)
            story.append(Spacer(1, 24))
        
        # Summary
        if (options.template.include_summary and 
            job.transcript_metadata and 
            job.transcript_metadata.summary):
            story.append(Paragraph("<b>Summary</b>", styles['Heading2']))
            story.append(Paragraph(job.transcript_metadata.summary, normal_style))
            story.append(Spacer(1, 12))
        
        # Keywords
        if (options.template.include_keywords and 
            job.transcript_metadata and 
            job.transcript_metadata.keywords):
            story.append(Paragraph("<b>Keywords</b>", styles['Heading2']))
            keywords_text = ", ".join(job.transcript_metadata.keywords)
            story.append(Paragraph(keywords_text, normal_style))
            story.append(Spacer(1, 12))
        
        # Transcript content
        story.append(Paragraph("<b>Transcript</b>", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        lines = job.transcript_content.strip().split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line.strip(), normal_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        pdf_io.seek(0)
        return pdf_io.getvalue()


# Service singleton
transcript_export_service = TranscriptExportService()